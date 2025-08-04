import streamlit as st
import os
import tempfile
import json
from datetime import datetime
from docx import Document
import re
from typing import Dict, Any, List
import requests
import io
import time
import PyPDF2
import fitz 
from pathlib import Path
import glob


class DocumentAnalyzer:
    """Classe para análisar documentos PDF e DOCX"""
    
    def __init__(self):
        self.supported_extensions = ['.pdf', '.docx']
    
    def extract_text_from_pdf(self, file_path: str) -> str:
        """Extrai texto de um arquivo PDF usando PyMuPDF"""
        try:
            doc = fitz.open(file_path)
            text = ""
            
            for page_num in range(doc.page_count):
                page = doc[page_num]
                text += page.get_text()
            
            doc.close()
            return text
            
        except Exception as e:
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    text = ""
                    
                    for page in pdf_reader.pages:
                        text += page.extract_text()
                    
                    return text
            except Exception as e2:
                raise Exception(f"Erro ao extrair texto do PDF: {str(e)} / {str(e2)}")
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """Extrai texto de um arquivo DOCX"""
        try:
            doc = Document(file_path)
            full_text = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text)
            
            return '\n'.join(full_text)
            
        except Exception as e:
            raise Exception(f"Erro ao extrair texto do DOCX: {str(e)}")
    
    def extract_text_from_file(self, file_path: str) -> str:
        """Extrai texto de acordo com a extensão do arquivo"""
        file_extension = Path(file_path).suffix.lower()
        
        if file_extension == '.pdf':
            return self.extract_text_from_pdf(file_path)
        elif file_extension == '.docx':
            return self.extract_text_from_docx(file_path)
        else:
            raise Exception(f"Tipo de arquivo não suportado: {file_extension}")
    
    def find_documents_in_folder(self, folder_path: str) -> List[str]:
        """Encontra todos os documentos suportados na pasta"""
        documents = []
        
        for extension in self.supported_extensions:
            pattern = os.path.join(folder_path, f"*{extension}")
            documents.extend(glob.glob(pattern))
        
        return sorted(documents)
    
    def analyze_single_document(self, file_path: str) -> Dict[str, Any]:
        """Analisa um único documento"""
        try:
            text_content = self.extract_text_from_file(file_path)
            
            if not text_content.strip():
                raise Exception("Documento vazio ou não foi possível extrair texto")
            
            text_content = re.sub(r'\s+', ' ', text_content).strip()
            
            sentences = re.split(r'[.!?]+(?:\s|$)', text_content)
            sentences = [s.strip() for s in sentences if s.strip() and len(s.strip()) > 3]
            
            paragraphs = [p.strip() for p in text_content.split('\n') if p.strip() and len(p.strip()) > 10]
            
            words = text_content.split()
            unique_words = set(word.lower().strip('.,!?;:"()[]') for word in words)
            
            complex_sentences = [s for s in sentences if len(s.split()) > 20]
            simple_sentences = [s for s in sentences if len(s.split()) < 10]
            
            exclamations = text_content.count('!')
            questions = text_content.count('?')
            semicolons = text_content.count(';')
            colons = text_content.count(':')
            
            transition_words = ['além disso', 'portanto', 'entretanto', 'assim', 'consequentemente', 
                              'por outro lado', 'dessa forma', 'contudo', 'no entanto', 'ademais',
                              'também', 'inclusive', 'sobretudo', 'principalmente', 'especialmente']
            transition_count = sum(text_content.lower().count(word) for word in transition_words)
            
            analysis = {
                'file_name': os.path.basename(file_path),
                'file_path': file_path,
                'total_paragraphs': len(paragraphs),
                'total_words': len(words),
                'total_sentences': len(sentences),
                'unique_words': len(unique_words),
                'avg_sentence_length': len(words) / max(1, len(sentences)),
                'avg_paragraph_length': len(words) / max(1, len(paragraphs)),
                'vocabulary_richness': len(unique_words) / max(1, len(words)),
                'complex_sentences': len(complex_sentences),
                'simple_sentences': len(simple_sentences),
                'exclamations': exclamations,
                'questions': questions,
                'semicolons': semicolons,
                'colons': colons,
                'transition_words_count': transition_count,
                'sample_paragraphs': paragraphs[:3] if paragraphs else [],
                'full_text': text_content,
                'sample_text': text_content[:1500] + "..." if len(text_content) > 1500 else text_content
            }
            
            return analysis
            
        except Exception as e:
            raise Exception(f"Erro ao analisar {os.path.basename(file_path)}: {str(e)}")
    
    def analyze_multiple_documents(self, folder_path: str, progress_callback=None) -> Dict[str, Any]:
        """Analisa múltiplos documentos e consolida o estilo"""
        documents = self.find_documents_in_folder(folder_path)
        
        if not documents:
            raise Exception(f"Nenhum documento PDF ou DOCX encontrado em: {folder_path}")
        
        st.info(f"Encontrados {len(documents)} documentos para análise")
        
        individual_analyses = []
        all_texts = []
        failed_files = []
        
        for i, doc_path in enumerate(documents):
            try:
                if progress_callback:
                    progress_callback(i / len(documents), f"Analisando: {os.path.basename(doc_path)}")
                
                analysis = self.analyze_single_document(doc_path)
                individual_analyses.append(analysis)
                all_texts.append(analysis['full_text'])
                
            except Exception as e:
                failed_files.append({'file': os.path.basename(doc_path), 'error': str(e)})
                st.warning(f"Falha ao analisar: {os.path.basename(doc_path)} - {str(e)}")
        
        if not individual_analyses:
            raise Exception("Nenhum documento foi analisado com sucesso")
        
        consolidated_text = '\n\n'.join(all_texts)
        
        total_words = sum(a['total_words'] for a in individual_analyses)
        total_sentences = sum(a['total_sentences'] for a in individual_analyses)
        total_paragraphs = sum(a['total_paragraphs'] for a in individual_analyses)
        total_unique_words = len(set(word.lower().strip('.,!?;:"()[]') 
                                   for text in all_texts 
                                   for word in text.split()))
        
        avg_sentence_length = sum(a['avg_sentence_length'] for a in individual_analyses) / len(individual_analyses)
        avg_vocabulary_richness = sum(a['vocabulary_richness'] for a in individual_analyses) / len(individual_analyses)
        
        consolidated_analysis = {
            'documents_count': len(individual_analyses),
            'failed_count': len(failed_files),
            'total_words': total_words,
            'total_sentences': total_sentences,
            'total_paragraphs': total_paragraphs,
            'unique_words_across_all': total_unique_words,
            'avg_sentence_length': avg_sentence_length,
            'avg_vocabulary_richness': avg_vocabulary_richness,
            'complex_sentences': sum(a['complex_sentences'] for a in individual_analyses),
            'simple_sentences': sum(a['simple_sentences'] for a in individual_analyses),
            'exclamations': sum(a['exclamations'] for a in individual_analyses),
            'questions': sum(a['questions'] for a in individual_analyses),
            'semicolons': sum(a['semicolons'] for a in individual_analyses),
            'colons': sum(a['colons'] for a in individual_analyses),
            'transition_words_count': sum(a['transition_words_count'] for a in individual_analyses),
            'sample_texts': [a['sample_text'] for a in individual_analyses[:5]],  # Primeiros 5 exemplos
            'full_consolidated_text': consolidated_text,
            'sample_consolidated_text': consolidated_text[:3000] + "..." if len(consolidated_text) > 3000 else consolidated_text,
            'individual_analyses': individual_analyses,
            'failed_files': failed_files,
            'document_names': [a['file_name'] for a in individual_analyses]
        }
        
        return consolidated_analysis


class OllamaTextGenerator:
    """Gerador de texto usando Ollama"""
    
    def __init__(self, model_name: str = "llama3.1"):
        self.model_name = model_name
        self.base_url = "http://localhost:11434"
        self.timeout = 120  
        
    def check_ollama_connection(self):
        """Verifica se o Ollama está rodando"""
        try:
            response = requests.get(f"{self.base_url}/api/tags", timeout=10)
            return response.status_code == 200
        except Exception as e:
            print(f"Erro de conexão com Ollama: {e}")
            return False
    
    def get_available_models(self):
        """Obtém lista de modelos instalados no Ollama"""
        try:
            response = requests.get(f"{self.base_url}/api/tags", timeout=10)
            if response.status_code == 200:
                data = response.json()
                return [model['name'] for model in data.get('models', [])]
            return []
        except Exception as e:
            print(f"Erro ao obter modelos: {e}")
            return []
    
    def generate_consolidated_style_analysis_prompt(self, analysis: Dict[str, Any]) -> str:
        """Cria prompt para análise de estilo consolidada de múltiplos documentos"""
        return f"""Analise esta coleção de {analysis['documents_count']} documentos brasileiros e descreva o estilo de escrita consolidado em português:

DADOS ESTATÍSTICOS CONSOLIDADOS:
- Total de documentos: {analysis['documents_count']}
- Total de palavras: {analysis['total_words']}
- Total de sentenças: {analysis['total_sentences']}
- Total de parágrafos: {analysis['total_paragraphs']}
- Vocabulário único: {analysis['unique_words_across_all']} palavras
- Média palavras/sentença: {analysis['avg_sentence_length']:.1f}
- Riqueza vocabular média: {analysis['avg_vocabulary_richness']:.3f}
- Sentenças complexas: {analysis['complex_sentences']}
- Pontuação especial: {analysis['exclamations']} exclamações, {analysis['questions']} perguntas

DOCUMENTOS ANALISADOS:
{', '.join(analysis['document_names'][:10])}{'...' if len(analysis['document_names']) > 10 else ''}

AMOSTRAS REPRESENTATIVAS DOS TEXTOS:
{analysis['sample_consolidated_text']}

Com base nesta análise consolidada, descreva em no minimo com 2200 palavras:
1. Tom geral e nível de formalidade predominante
2. Padrões estruturais das frases e organização textual
3. Características do vocabulário e terminologia típica
4. Estilo de conectivos e transições entre ideias
5. Particularidades distintivas deste conjunto de documentos

Seja específico sobre os padrões consistentes encontrados across todos os documentos."""
    
    def generate_content_prompt(self, style_description: str, topic: str) -> str:
        """Cria prompt para geração baseada no estilo consolidado"""
        return f"""Escreva um texto em português sobre "{topic}" seguindo exatamente o estilo consolidado identificado:

ESTILO CONSOLIDADO IDENTIFICADO:
{style_description}

INSTRUÇÕES ESPECÍFICAS:
- Mantenha absoluta fidelidade ao tom e formalidade identificados
- Use as mesmas estruturas frasais e padrões organizacionais
- Aplique o vocabulário e terminologia no mesmo nível
- Utilize conectivos e transições no estilo identificado
- Inclua título apropriado ao estilo
- Mantenha consistência com os padrões dos documentos originais

TEMA ESPECÍFICO: {topic}

Produza apenas o texto final completo, com 2200 palavras no minimo."""
    
    def call_ollama_api(self, prompt: str, max_tokens: int = 2000, temperature: float = 0.7) -> str:
        """Chama a API do Ollama"""
        try:
            payload = {
                "model": self.model_name,
                "prompt": prompt,
                "stream": False,
                "options": {
                    "temperature": temperature,
                    "top_p": 0.9,
                    "top_k": 40,
                    "num_predict": max_tokens,
                    "stop": ["Human:", "Assistant:", "\n\n\n"]
                }
            }
            
            response = requests.post(
                f"{self.base_url}/api/generate",
                json=payload,
                timeout=self.timeout
            )
            
            if response.status_code == 200:
                result = response.json()
                return result.get('response', '').strip()
            else:
                raise Exception(f"Erro HTTP {response.status_code}: {response.text}")
                
        except requests.exceptions.Timeout:
            raise Exception(f"Timeout após {self.timeout} segundos. Tente um modelo mais rápido.")
        except requests.exceptions.ConnectionError:
            raise Exception("Erro de conexão com Ollama. Verifique se está rodando.")
        except Exception as e:
            raise Exception(f"Erro na API Ollama: {str(e)}")
    
    def analyze_consolidated_style(self, consolidated_analysis: Dict[str, Any]) -> str:
        """Analisa o estilo consolidado usando Ollama"""
        prompt = self.generate_consolidated_style_analysis_prompt(consolidated_analysis)
        return self.call_ollama_api(prompt, max_tokens=1200, temperature=0.5)
    
    def generate_content(self, style_description: str, topic: str) -> str:
        """Gera conteúdo baseado no estilo consolidado"""
        prompt = self.generate_content_prompt(style_description, topic)
        return self.call_ollama_api(prompt, max_tokens=2500, temperature=0.7)


class DocumentGenerator:
    """Classe para gerar documentos DOCX"""
    
    def save_to_docx(self, content: str, output_path: str) -> bool:
        """Salva conteúdo em arquivo DOCX"""
        try:
            doc = Document()
            
            lines = content.split('\n')
            
            for line in lines:
                line = line.strip()
                if line:
                    if line.startswith('#'):
                        title_text = line.replace('#', '').strip()
                        level = 1 if line.startswith('# ') else 2
                        doc.add_heading(title_text, level=level)
                    elif line.isupper() and len(line.split()) <= 8 and len(line) > 3:
                        doc.add_heading(line, level=1)
                    elif line.endswith(':') and len(line.split()) <= 6 and len(line) > 5:
                        doc.add_heading(line.replace(':', ''), level=2)
                    else:
                        doc.add_paragraph(line)
            
            doc.save(output_path)
            return True
            
        except Exception as e:
            print(f"Erro ao salvar documento: {str(e)}")
            return False


def init_session_state():
    """Inicializa o estado da sessão do Streamlit"""
    if 'messages' not in st.session_state:
        st.session_state.messages = []
    if 'documents_analyzed' not in st.session_state:
        st.session_state.documents_analyzed = []
    if 'consolidated_analysis' not in st.session_state:
        st.session_state.consolidated_analysis = None
    if 'style_description' not in st.session_state:
        st.session_state.style_description = None
    if 'current_folder' not in st.session_state:
        st.session_state.current_folder = None


def main():
    """Função principal da aplicação Streamlit"""
    
    st.set_page_config(
        page_title="IA Agent Procedure | Powered by WOOD",
        page_icon="📚",
        layout="wide",
        initial_sidebar_state="expanded"
    )
    
    init_session_state()
    
    st.title("SmartOps AI")
    st.markdown("*Sistema inteligente para análise de estilo consolidado de múltiplos documentos PDF*")
    
    with st.sidebar:
        st.header("⚙️ Configurações")
        
        st.subheader("Configuração Ollama")
        
        text_generator = OllamaTextGenerator()
        
        connection_status = text_generator.check_ollama_connection()
        
        if connection_status:
            st.success("✅ Ollama conectado!")
            available_models = text_generator.get_available_models()
            
            if available_models:
                preferred_models = [m for m in available_models if any(x in m.lower() for x in ['llama3.1', 'llama3', 'mistral', 'phi'])]
                other_models = [m for m in available_models if m not in preferred_models]
                sorted_models = preferred_models + other_models
                
                selected_model = st.selectbox(
                    "Modelo:",
                    sorted_models,
                    index=0,
                    help="Para múltiplos documentos, prefira modelos maiores"
                )
                text_generator.model_name = selected_model
                
                timeout_options = {60: "Normal (60s)", 120: "Lento (120s)", 180: "Muito Lento (180s)"}
                selected_timeout = st.selectbox(
                    "Timeout:",
                    list(timeout_options.keys()),
                    index=1,
                    format_func=lambda x: timeout_options[x]
                )
                text_generator.timeout = selected_timeout
                
                st.info(f"📊 Modelo: {selected_model}")
            else:
                st.warning("⚠️ Nenhum modelo disponível")
                st.markdown("Execute: `ollama pull llama3.1`")
                st.stop()
        else:
            st.error("❌ Ollama não está rodando!")
            st.markdown("""
            **Como iniciar o Ollama:**
            1. Abra terminal/cmd
            2. Execute: `ollama serve`
            3. Em outro terminal: `ollama pull llama3.1`
            """)
            st.stop()
        
        st.divider()
        
        st.subheader("🔧 Opções de Análise")
        show_detailed_analysis = st.checkbox("Análise detalhada por documento", value=False)
        show_consolidated_stats = st.checkbox("Estatísticas consolidadas", value=True)
        show_style_analysis = st.checkbox("Descrição do estilo", value=True)
        
        st.divider()
        
        st.subheader(" Pasta Atual")
        current_path = os.getcwd()
        st.code(current_path)
        
        analyzer = DocumentAnalyzer()
        docs_in_folder = analyzer.find_documents_in_folder(current_path)
        
        if docs_in_folder:
            st.success(f"✅ {len(docs_in_folder)} documentos encontrados")
            
            with st.expander("📄 Documentos encontrados"):
                for doc in docs_in_folder:
                    file_size = os.path.getsize(doc) / 1024  # KB
                    st.write(f"• {os.path.basename(doc)} ({file_size:.1f} KB)")
        else:
            st.warning("⚠️ Nenhum PDF/DOCX encontrado na pasta raiz")
    
    # Conteúdo principal
    col1, col2 = st.columns([1, 1])
    
    with col1:
        
        if docs_in_folder:
            st.info(f" {len(docs_in_folder)} documentos prontos para análise de escrita")
                        
            new_topic = st.text_input(
                "💡 Tema para o novo texto:",
                placeholder="Ex: Procedimento para uso de IA em operações",
                help="Digite o tema sobre o qual você quer gerar um novo texto"
            )
            
            if st.button("🚀 Analisar Documentos e Gerar", type="primary", use_container_width=True):
                if new_topic.strip():
                    try:
                        start_time = time.time()
                        
                        analyzer = DocumentAnalyzer()
                        doc_generator = DocumentGenerator()
                        
                        
                        with st.spinner("📊 Analisando múltiplos documentos..."):
                            progress_container = st.container()
                            progress_bar = progress_container.progress(0)
                            status_text = progress_container.empty()
                            
                            def progress_callback(progress, message):
                                progress_bar.progress(progress)
                                status_text.text(message)
                            
                            # Correção: Adicionar current_path como primeiro argumento
                            consolidated_analysis = analyzer.analyze_multiple_documents(
                                current_path,  # Adicionar o caminho da pasta
                                progress_callback
                            )
                            
                            progress_bar.progress(1.0)
                            status_text.text("✅ Análise consolidada concluída!")
                        
                        st.success(f"✅ {consolidated_analysis['documents_count']} documentos analisados!")
                        
                        if consolidated_analysis['failed_count'] > 0:
                            st.warning(f"⚠️ {consolidated_analysis['failed_count']} documentos falharam na análise")
                        
                        if show_consolidated_stats:
                            with st.expander("📊 Estatísticas Consolidadas", expanded=True):
                                col_a, col_b, col_c, col_d = st.columns(4)
                                with col_a:
                                    st.metric("Documentos", consolidated_analysis['documents_count'])
                                    st.metric("Total Palavras", f"{consolidated_analysis['total_words']:,}")
                                with col_b:
                                    st.metric("Total Sentenças", f"{consolidated_analysis['total_sentences']:,}")
                                    st.metric("Total Parágrafos", f"{consolidated_analysis['total_paragraphs']:,}")
                                with col_c:
                                    st.metric("Vocabulário Único", f"{consolidated_analysis['unique_words_across_all']:,}")
                                    st.metric("Média Sent./Doc", f"{consolidated_analysis['total_sentences']/consolidated_analysis['documents_count']:.0f}")
                                with col_d:
                                    st.metric("Média Palavras/Sent.", f"{consolidated_analysis['avg_sentence_length']:.1f}")
                                    st.metric("Riqueza Vocabular", f"{consolidated_analysis['avg_vocabulary_richness']:.3f}")
                        
                        if show_detailed_analysis:
                            with st.expander("📋 Análise Detalhada por Documento", expanded=False):
                                for analysis in consolidated_analysis['individual_analyses']:
                                    st.write(f"**{analysis['file_name']}**")
                                    col_x, col_y, col_z = st.columns(3)
                                    with col_x:
                                        st.write(f"Palavras: {analysis['total_words']:,}")
                                    with col_y:
                                        st.write(f"Sentenças: {analysis['total_sentences']:,}")
                                    with col_z:
                                        st.write(f"Parágrafos: {analysis['total_paragraphs']:,}")
                                    st.divider()
                        
                        with st.spinner("🎨 Analisando estilo consolidado..."):
                            style_description = text_generator.analyze_consolidated_style(consolidated_analysis)
                            st.session_state.style_description = style_description
                        
                        st.success("✅ Estilo consolidado analisado!")
                        
                        if show_style_analysis:
                            with st.expander("🎨 Estilo Consolidado Identificado", expanded=True):
                                st.markdown(style_description)
                        
                        with st.spinner("✍️ Gerando novo texto baseado no estilo ..."):
                            generated_content = text_generator.generate_content(style_description, new_topic)
                        
                        st.success("✅ Texto gerado com base no estilo de escrita!")
                        
                        with st.spinner("Salvando documento..."):
                            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                            output_filename = f"texto_consolidado_{timestamp}.docx"
                            temp_output_path = os.path.join(tempfile.gettempdir(), output_filename)
                            
                            success = doc_generator.save_to_docx(generated_content, temp_output_path)
                            
                            if success:
                                st.success("✅ Documento salvo!")
                            else:
                                st.error("❌ Erro ao salvar documento")
                        
                        word_count = len(generated_content.split())
                        processing_time = time.time() - start_time
                        
                        st.session_state.documents_analyzed.append({
                            'type': 'multi_document_analysis',
                            'documents_count': consolidated_analysis['documents_count'],
                            'topic': new_topic,
                            'timestamp': datetime.now().strftime("%H:%M:%S"),
                            'status': 'Concluído',
                            'output_path': temp_output_path,
                            'output_filename': output_filename,
                            'word_count': word_count,
                            'processing_time': f"{processing_time:.1f}s"
                        })
                        
                        st.session_state.consolidated_analysis = {
                            'content': generated_content,
                            'output_path': temp_output_path,
                            'output_filename': output_filename,
                            'word_count': word_count,
                            'processing_time': processing_time,
                            'documents_analyzed': consolidated_analysis['documents_count'],
                            'style_description': style_description
                        }
                        
                        st.balloons()
                        
                    except Exception as e:
                        st.error(f"❌ Erro durante o processamento: {str(e)}")
                        
                        with st.expander("🔍 Detalhes do Erro"):
                            st.code(str(e))
                            st.info("Verifique se os PDFs não estão corrompidos ou protegidos por senha")
                else:
                    st.warning("Digite um tema para o novo texto")
        else:
            st.warning("⚠️ Nenhum documento PDF ou DOCX encontrado na pasta raiz")
            st.info("Adicione arquivos PDF ou DOCX na pasta do seu código e recarregue a página")
    
    with col2:
        if st.session_state.consolidated_analysis:
            st.header("📄 Resultado")
            
            col_info1, col_info2, col_info3 = st.columns(3)
            with col_info1:
                st.metric(" Docs Analisados", st.session_state.consolidated_analysis['documents_analyzed'])
            with col_info2:
                st.metric("Palavras", st.session_state.consolidated_analysis['word_count'])
            with col_info3:
                st.metric("Tempo", f"{st.session_state.consolidated_analysis['processing_time']:.1f}s")
            
            if os.path.exists(st.session_state.consolidated_analysis['output_path']):
                with open(st.session_state.consolidated_analysis['output_path'], 'rb') as file:
                    st.download_button(
                        label="Baixar DOCX",
                        data=file.read(),
                        file_name=st.session_state.consolidated_analysis['output_filename'],
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
            
            with st.expander("Resumo do Estilo", expanded=False):
                st.markdown(st.session_state.consolidated_analysis['style_description'])
            
            with st.expander("Prévia do Texto Gerado", expanded=True):
                st.text_area(
                    "Conteúdo:",
                    value=st.session_state.consolidated_analysis['content'],
                    height=400,
                    disabled=True
                )
        else:
            
            if docs_in_folder:
                st.success("📋 Documentos detectados - pronto para começar!")
            else:
                st.info("📋 Adicione documentos PDF na pasta raiz para começar!")
        
        # Histórico de análises
        if st.session_state.documents_analyzed:
            st.divider()
            st.subheader("📚 Histórico de Análises")
            
            for i, analysis in enumerate(reversed(st.session_state.documents_analyzed)):
                with st.expander(f"📊 Análise #{len(st.session_state.documents_analyzed)-i}"):
                    if analysis.get('type') == 'multi_document_analysis':
                        st.write(f"**Tipo:** Análise Multi-Documentos")
                        st.write(f"**Documentos:** {analysis['documents_count']} arquivos")
                    else:
                        st.write(f"**Tipo:** Análise Individual")
                    
                    st.write(f"**Tema:** {analysis['topic']}")
                    st.write(f"**Horário:** {analysis['timestamp']}")
                    st.write(f"**Palavras:** {analysis['word_count']}")
                    st.write(f"**Tempo:** {analysis['processing_time']}")
                    st.write(f"**Status:** {analysis['status']}")
    
    try:
        temp_dir = tempfile.gettempdir()
        for file in os.listdir(temp_dir):
            if file.startswith('texto_consolidado_') and file.endswith('.docx'):
                file_path = os.path.join(temp_dir, file)
                if os.path.getctime(file_path) < time.time() - 3600:
                    os.remove(file_path)
    except:
        pass
    
    st.divider()
    st.markdown("""
    <div style='text-align: center; color: #666; font-size: 0.8em;'>
         SmartOps AI| Powered by WOOD<br>
    </div>
    """, unsafe_allow_html=True)


if __name__ == "__main__":
    main()